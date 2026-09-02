
from dotenv import load_dotenv
import streamlit as st
import docx

# ReportLab imports for PDF generation matching original design

from chain import cv_chain
from docs_builder import create_docx_cv
from pdf_builder import create_pdf_cv
from schemas import TailoredResume

load_dotenv()

# ==============================================================================
# 4. Streamlit User Interface
# ==============================================================================

def extract_text_from_docx(file_bytes) -> str:
    """Extracts all text from a Word (.docx) file stream."""
    doc = docx.Document(file_bytes)
    full_text = []

    # 1. Extract regular paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    # 2. Extract text inside tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text.strip())

    return "\n".join(full_text)

st.set_page_config(page_title="AI Resume Tailor", page_icon="📄", layout="wide")
st.title("AI Resume Tailoring Agent 📄")
st.caption("Tailor your master resume with strict role integrity and export directly to PDF.")

col1, col2 = st.columns(2)

with col1:
    st.subheader("Base Resume")
    # File uploader for Word document
    uploaded_file = st.file_uploader(
        "Upload your master resume (.docx):",
        type=["docx"],
        help="Upload your Word CV file to extract background data automatically."
    )

with col2:
    st.subheader("Job Description")
    job_desc_input = st.text_area(
        "Paste target job description:",
        height=320,
        placeholder="Paste target job responsibilities and requirements..."
    )

if st.button("Optimize & Generate Tailored Resume 🎯", type="primary"):
    if not uploaded_file or not job_desc_input.strip():
        st.warning("Please upload your base CV (.docx) and provide the target job description.")
    else:
        with st.spinner("Analyzing alignment and compiling documents..."):
            try:
                base_cv_input = extract_text_from_docx(uploaded_file)
                tailored_result: TailoredResume = cv_chain.invoke({
                    "base_cv": base_cv_input,
                    "job_description": job_desc_input
                })

                # to prevent re-processing on rerun, store results in session state
                st.session_state["tailored_result"] = tailored_result
                st.session_state["pdf_bytes"] = create_pdf_cv(tailored_result).getvalue()
                st.session_state["docx_bytes"] = create_docx_cv(tailored_result).getvalue()
                st.session_state["safe_name"] = tailored_result.personal_info.name.strip().replace(" ", "_")

            except Exception as e:
                st.error(f"Execution Error: {e}")

# show download buttons and JSON preview only if tailored_result exists in session state
if "tailored_result" in st.session_state:
    st.success("Resume tailored successfully!")
    
    dl_col1, dl_col2 = st.columns(2)
    safe_name = st.session_state["safe_name"]

    with dl_col1:
        st.download_button(
            label="Download as PDF 📄",
            data=st.session_state["pdf_bytes"],
            file_name=f"{safe_name}_Tailored_CV.pdf",
            mime="application/pdf",
            use_container_width=True,
            key="dl_pdf_btn"
        )

    with dl_col2:
        st.download_button(
            label="Download as Word 📝 (.docx)",
            data=st.session_state["docx_bytes"],
            file_name=f"{safe_name}_Tailored_CV.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="dl_docx_btn"
        )

    with st.expander("Preview Extracted Structured Resume (JSON)"):
        st.json(st.session_state["tailored_result"].model_dump())