import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from schemas import TailoredResume

def create_docx_cv(data: TailoredResume) -> io.BytesIO:
    doc = Document()

    # define a consistent accent color for headings
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

    # define a consistent accent color for headings
    ACCENT_COLOR = RGBColor(0xF3, 0x72, 0xAE)

    # title and name
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.paragraph_format.space_after = Pt(2)
    header_p.paragraph_format.space_before = Pt(0)
    
    name_run = header_p.add_run(f"{data.personal_info.name} - {data.personal_info.title}")
    name_run.bold = True
    name_run.font.size = Pt(14)
    name_run.font.color.rgb = ACCENT_COLOR

    # contact information
    c = data.personal_info.contact
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(6)
    contact_p.paragraph_format.space_before = Pt(0)

    line1 = []
    if c.phone: line1.append(f"Phone: {c.phone}")
    if c.linkedin: line1.append(f"LinkedIn: {c.linkedin}")
    line2 = []
    if c.email: line2.append(f"E-mail: {c.email}")
    if c.github: line2.append(f"GitHub: {c.github}")

    run_contact = contact_p.add_run(" | ".join(line1) + "\n" + " | ".join(line2))
    run_contact.font.size = Pt(8.5)

    # helper function to add section headings
    def add_docx_heading(title: str):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(6)
        h.paragraph_format.space_after = Pt(2)
        run = h.add_run(title)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = ACCENT_COLOR

    # ABOUT ME
    add_docx_heading("ABOUT ME")
    p_about = doc.add_paragraph(data.about_me)
    p_about.paragraph_format.space_after = Pt(4)
    p_about.paragraph_format.space_before = Pt(0)
    for run in p_about.runs: run.font.size = Pt(8.5)

    # WORK EXPERIENCE
    add_docx_heading("WORK EXPERIENCE")
    for job in data.work_experience:
        p_role = doc.add_paragraph()
        p_role.paragraph_format.space_before = Pt(3)
        p_role.paragraph_format.space_after = Pt(1)
        r_title = p_role.add_run(f"{job.role}, {job.organization}")
        r_title.bold = True
        r_title.font.size = Pt(8.5)
        r_dates = p_role.add_run(f" ({job.period})")
        r_dates.font.size = Pt(8.5)

        for bullet in job.highlights:
            p_bullet = doc.add_paragraph(bullet, style='List Bullet')
            p_bullet.paragraph_format.space_after = Pt(1)
            p_bullet.paragraph_format.space_before = Pt(0)
            for run in p_bullet.runs: run.font.size = Pt(8.5)

    # TECHNICAL PROJECTS
    if data.technical_projects:
        add_docx_heading("TECHNICAL PROJECTS")
        for proj in data.technical_projects:
            p_proj = doc.add_paragraph()
            p_proj.paragraph_format.space_before = Pt(3)
            p_proj.paragraph_format.space_after = Pt(1)
            r_name = p_proj.add_run(f"{proj.name} | ")
            r_name.bold = True
            r_name.font.size = Pt(8.5)
            r_tech = p_proj.add_run(proj.technologies)
            r_tech.italic = True
            r_tech.font.size = Pt(8.5)

            for bullet in proj.highlights:
                p_bullet = doc.add_paragraph(bullet, style='List Bullet')
                p_bullet.paragraph_format.space_after = Pt(1)
                p_bullet.paragraph_format.space_before = Pt(0)
                for run in p_bullet.runs: run.font.size = Pt(8.5)

    # PROFESSIONAL SKILLS
    add_docx_heading("PROFESSIONAL SKILLS")
    s = data.professional_skills
    skills_map = [
        ("Languages", s.languages),
        ("Software Development", s.software_development),
        ("Data Analysis & Visualization", s.data_analysis_and_visualization),
        ("Environments", s.environments),
        ("Bioinformatics Tools", s.bioinformatics_and_domain_tools),
        ("Machine Learning", s.machine_learning),
        ("Soft Skills", s.soft_skills),
    ]
    for label, items in skills_map:
        if items:
            p_skill = doc.add_paragraph()
            p_skill.paragraph_format.space_after = Pt(1)
            p_skill.paragraph_format.space_before = Pt(0)
            r_lbl = p_skill.add_run(f"{label}: ")
            r_lbl.bold = True
            r_lbl.font.size = Pt(8.5)
            r_val = p_skill.add_run(f"{', '.join(items)}.")
            r_val.font.size = Pt(8.5)

    # EDUCATION
    add_docx_heading("EDUCATION")
    for edu in data.education:
        p_edu = doc.add_paragraph()
        p_edu.paragraph_format.space_after = Pt(1)
        p_edu.paragraph_format.space_before = Pt(0)
        r_deg = p_edu.add_run(f"{edu.degree}, {edu.institution} ({edu.period})")
        r_deg.bold = True
        r_deg.font.size = Pt(8.5)
        if edu.honors_or_notes:
            r_notes = p_edu.add_run(f", {edu.honors_or_notes}")
            r_notes.font.size = Pt(8.5)

    # MILITARY SERVICE
    if data.military_service:
        mil = data.military_service[0] if isinstance(data.military_service, list) and data.military_service else data.military_service
        if getattr(mil, 'role', None):
            add_docx_heading("MILITARY SERVICE")
            p_mil = doc.add_paragraph()
            p_mil.paragraph_format.space_after = Pt(1)
            p_mil.paragraph_format.space_before = Pt(0)
            r_mil = p_mil.add_run(f"{mil.role} ({mil.period})")
            r_mil.bold = True
            r_mil.font.size = Pt(8.5)
            p_desc = doc.add_paragraph(mil.description)
            p_desc.paragraph_format.space_after = Pt(2)
            p_desc.paragraph_format.space_before = Pt(0)
            for run in p_desc.runs: run.font.size = Pt(8.5)

    # LANGUAGES
    if data.languages:
        add_docx_heading("LANGUAGES")
        p_lang = doc.add_paragraph()
        p_lang.paragraph_format.space_after = Pt(2)
        p_lang.paragraph_format.space_before = Pt(0)
        parts = [f"{l.language}: {l.proficiency}" for l in data.languages]
        r_lang = p_lang.add_run(" | ".join(parts))
        r_lang.font.size = Pt(8.5)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer